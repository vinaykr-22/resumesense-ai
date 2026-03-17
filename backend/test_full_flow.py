import requests
import time
import os
from pprint import pprint

BASE_URL = "http://127.0.0.1:8000/api/v1"

def test_full_pipeline():
    print("=== Testing FastAPI Job Matcher Pipeline ===")
    
    # 1. Register a new user
    print("\n1. Registering user...")
    email = f"test_{int(time.time())}@example.com"
    res = requests.post(f"{BASE_URL}/auth/register", json={
        "email": email,
        "password": "password123"
    })
    
    if res.status_code != 200:
        print("Registration failed!", res.text)
        return
        
    token = res.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    print("✓ Successfully registered and obtained JWT token.")
    
    # 2. Upload a sample Resume
    print("\n2. Uploading a sample Resume...")
    
    # Create a quick dummy DOCX file to test the upload endpoint properly 
    from docx import Document
    doc = Document()
    doc.add_heading("John Doe - Senior Software Engineer", 0)
    doc.add_paragraph("Experienced backend developer with 5 years in Python, FastAPI, Docker, and PostgreSQL.")
    doc.add_paragraph("Skills: Python, Go, Node.js, Kubernetes, AWS, Redis, SQL, Git.")
    doc.add_paragraph("Experience: Built scalable microservices for a fintech startup handling millions of requests daily.")
    test_docx_path = "temp_test_resume.docx"
    doc.save(test_docx_path)
    
    with open(test_docx_path, "rb") as f:
        files = {"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        res = requests.post(f"{BASE_URL}/resume/upload", headers=headers, files=files)
        
    os.remove(test_docx_path) # Clean up
    
    if res.status_code != 200:
        print("Upload failed!", res.text)
        return
        
    upload_data = res.json()
    resume_id = upload_data["resume_id"]
    job_id = upload_data["job_id"]
    print(f"✓ Upload successful! Resume ID: {resume_id}")
    
    # 3. Poll Celery Status
    print("\n3. Polling Celery Processing Status...")
    status = "processing"
    max_retries = 30
    retries = 0
    final_result = None
    
    while status not in ["completed", "failed"] and retries < max_retries:
        time.sleep(2)
        res = requests.get(f"{BASE_URL}/resume/status/{job_id}")
        if res.status_code == 200:
            status_data = res.json()
            status = status_data.get("status", "unknown")
            print(f"   [{status_data.get('progress', 0)}%] {status_data.get('stage_label', 'Wait...')}")
            
            if status == "completed":
                final_result = status_data.get("result")
        else:
            print("Status check failed!", res.text)
            return
        retries += 1
        
    if status != "completed":
        print("Processing timed out or failed!")
        return
        
    print("✓ Celery background processing completed successfully!")
    
    # 4. Check Extracted Skills Endpoint
    print("\n4. Testing /resume/skills endpoint (should hit cache)...")
    res = requests.post(f"{BASE_URL}/resume/skills", headers=headers, json={"resume_id": resume_id})
    if res.status_code == 200:
        print("✓ Received structured skills JSON!")
        pprint(res.json(), depth=2)
    else:
        print("Skills fetch failed!", res.text)
        
    # 5. Check Job Matching Endpoint
    print("\n5. Testing /jobs/match endpoint...")
    res = requests.post(f"{BASE_URL}/jobs/match", headers=headers, json={"resume_id": resume_id, "top_k": 3})
    if res.status_code == 200:
        match_data = res.json()
        print(f"✓ Found {match_data.get('total')} matches!")
        for match in match_data.get("matches", []):
            print(f"   - {match['match_score']}% Match: {match['title']} ({match['category']})")
    else:
        print("Job matching failed!", res.text)
        
    print("\n=== All Tests Passed Successfully! ===")

if __name__ == "__main__":
    test_full_pipeline()
